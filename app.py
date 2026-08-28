import io
import os
import re
import time
import random
import sqlite3
from datetime import datetime
from urllib.parse import quote_plus

import pandas as pd
import requests
import streamlit as st
from docx import Document
from google import genai
from groq import Groq
from pypdf import PdfReader


# ==========================================
# 1. إعدادات الصفحة والتصميم
# ==========================================
st.set_page_config(
    page_title="محطة عمل الباحث الذكي - ANRN",
    page_icon="🏛️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;500;700;800&display=swap');

    .stApp, .stMarkdown, p, h1, h2, h3, h4, label, input, textarea, select {
        font-family: 'Tajawal', sans-serif !important;
        direction: rtl !important;
        text-align: right !important;
    }

    @media (max-width: 768px) {
        .block-container { padding: 1rem 0.5rem !important; }
        h1 { font-size: 1.4rem !important; }
        h2 { font-size: 1.15rem !important; }
        .stButton>button { font-size: 0.85rem !important; padding: 0.5rem !important; }
    }

    .stButton>button {
        width: 100%;
        border-radius: 12px;
        font-weight: bold;
    }
</style>
""",
    unsafe_allow_html=True,
)


# ==========================================
# 2. أسرار التطبيق والجلسة
# ==========================================
def get_secret(name: str, default: str = "") -> str:
    """قراءة السر من Streamlit Secrets ثم من متغيرات البيئة."""
    try:
        value = st.secrets.get(name, default)
        if value:
            return str(value)
    except Exception:
        pass
    return os.getenv(name, default)


DEFAULT_GROQ_KEY = get_secret("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")
DEFAULT_GEMINI_KEY = get_secret("AQ.Ab8RN6JI7XuW1iL9Iy1mvC-eTpI1je3WDSB1A9Q1nlpJJylNUQ")
DEFAULT_S2_KEY = get_secret("SEMANTIC_SCHOLAR_API_KEY")
ADMIN_PIN = get_secret("Maghnia.2026")
DB_PATH = get_secret("DB_PATH", "smart_researcher_logs.db")

if "form_title" not in st.session_state:
    st.session_state["form_title"] = (
        "أثر تطبيقات الذكاء الاصطناعي على جودة التعليم العالي والحوكمة الأكاديمية"
    )
if "form_field" not in st.session_state:
    st.session_state["form_field"] = "علوم التسيير - إدارة المنظمات"


def reset_all_fields():
    st.session_state["form_title"] = ""
    st.session_state["form_field"] = ""
    st.session_state.pop("results", None)
    st.session_state.pop("current_log_id", None)
    st.rerun()


# ==========================================
# 3. قاعدة البيانات
# ==========================================
def get_db_connection():
    return sqlite3.connect(DB_PATH, timeout=20)


def init_db():
    with get_db_connection() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS research_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT,
                researcher_name TEXT,
                role TEXT,
                affiliation TEXT,
                email TEXT,
                phone TEXT,
                title TEXT,
                field TEXT,
                language TEXT,
                feedback TEXT
            )
            """
        )
        conn.commit()


def save_research_log(data):
    try:
        with get_db_connection() as conn:
            cur = conn.execute(
                """
                INSERT INTO research_logs (
                    timestamp, researcher_name, role, affiliation, email, phone,
                    title, field, language, feedback
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    data.get("name", ""),
                    data.get("role", ""),
                    data.get("affiliation", ""),
                    data.get("email", ""),
                    data.get("phone", ""),
                    data.get("title", ""),
                    data.get("field", ""),
                    data.get("language", ""),
                    "",
                ),
            )
            conn.commit()
            return cur.lastrowid
    except Exception as exc:
        st.warning(f"تعذر حفظ سجل البحث محلياً: {exc}")
        return None


def update_feedback(log_id, feedback_text):
    try:
        with get_db_connection() as conn:
            conn.execute(
                "UPDATE research_logs SET feedback = ? WHERE id = ?",
                (feedback_text, log_id),
            )
            conn.commit()
        return True
    except Exception as exc:
        st.warning(f"تعذر حفظ الملاحظة: {exc}")
        return False


init_db()


# ==========================================
# 4. أدوات مساعدة للبحث
# ==========================================
HTTP_HEADERS = {
    "User-Agent": "ANRN-Smart-Researcher/1.0 (academic research application)"
}


def safe_get_json(url, *, params=None, headers=None, timeout=15):
    merged_headers = {**HTTP_HEADERS, **(headers or {})}
    response = requests.get(
        url,
        params=params,
        headers=merged_headers,
        timeout=timeout,
    )
    response.raise_for_status()
    return response.json()


def first_author_openalex(item):
    authorships = item.get("authorships") or []
    if authorships:
        author = (authorships[0].get("author") or {}).get("display_name")
        if author:
            return author
    return "Auteur non renseigné"


def first_author_semantic(item):
    authors = item.get("authors") or []
    if authors and authors[0].get("name"):
        return authors[0]["name"]
    return "Auteur non renseigné"


def crossref_year(item):
    for key in ("published-print", "published-online", "issued", "created"):
        block = item.get(key) or {}
        parts = block.get("date-parts") or []
        if parts and parts[0]:
            return parts[0][0]
    return "N/A"


def crossref_author(item):
    authors = item.get("author") or []
    if not authors:
        return "Auteur non renseigné"
    first = authors[0]
    name = " ".join(
        part for part in [first.get("given", ""), first.get("family", "")] if part
    ).strip()
    return name or "Auteur non renseigné"


def epmc_author(item):
    author_string = item.get("authorString")
    if author_string:
        return author_string.split(",")[0].strip()
    return "Auteur non renseigné"


def deduplicate_papers(papers):
    seen = set()
    unique = []
    for paper in papers:
        key = (
            str(paper.get("doi") or "").strip().lower()
            or str(paper.get("title") or "").strip().lower()
        )
        if not key or key in seen:
            continue
        seen.add(key)
        unique.append(paper)
    return unique


def extract_english_keywords(title):
    keywords = []
    if any(w in title for w in ["تحول رقمي", "رقمنة", "تكنولوجيا", "رقمي"]):
        keywords.append("digital transformation")
    if any(w in title for w in ["ذكاء اصطناعي", "توليدي", "خوارزم"]):
        keywords.append("artificial intelligence")
    if any(w in title for w in ["حوكمة", "إدارة", "تسيير", "تنظيم"]):
        keywords.append("governance")
    if any(w in title for w in ["جامع", "تعليم عالي", "أكاديمي"]):
        keywords.append("higher education")
    if any(w in title for w in ["جودة", "أداء", "نضج"]):
        keywords.append("quality management")
    return " ".join(dict.fromkeys(keywords)) if keywords else title


def crawl_academic_papers(query, platforms, semantic_key=""):
    """
    يرجع:
      papers       = مقالات حقيقية مسترجعة عبر APIs عامة
      portal_links = روابط منصات لا نملك لها API عامة موثقة في هذا التطبيق
      errors       = أخطاء الشبكة/API لتسهيل التشخيص
    """
    papers = []
    portal_links = []
    errors = []

    en_query = extract_english_keywords(query)
    en_for_web = quote_plus(en_query)

    # 1) ASJP
    # لا نضيف مرجعاً وهمياً. ASJP توفر بوابة بحث متقدم، لذلك نعرض رابطها للمستخدم.
    if any("ASJP" in p for p in platforms):
        portal_links.append(
            {
                "name": "ASJP - Recherche avancée",
                "url": "https://asjp.cerist.dz/en/advancedResearch",
                "note": f"ابحث داخل ASJP بالكلمات: {query}",
            }
        )

    # 2) OpenAlex
    if any("OpenAlex" in p for p in platforms):
        try:
            data = safe_get_json(
                "https://api.openalex.org/works",
                params={
                    "search": en_query,
                    "per-page": 5,
                    "sort": "cited_by_count:desc",
                },
            )
            for item in data.get("results") or []:
                raw_doi = item.get("doi") or ""
                clean_doi = raw_doi.replace("https://doi.org/", "").strip()
                work_id = item.get("id") or ""
                papers.append(
                    {
                        "title": item.get("title") or "Sans titre",
                        "author": first_author_openalex(item),
                        "year": item.get("publication_year") or "N/A",
                        "doi": clean_doi,
                        "url": raw_doi or work_id,
                        "source": "OpenAlex",
                    }
                )
        except Exception as exc:
            errors.append(f"OpenAlex: {exc}")

    # 3) Semantic Scholar
    if any("Semantic Scholar" in p for p in platforms):
        try:
            headers = {"x-api-key": semantic_key} if semantic_key else {}
            data = safe_get_json(
                "https://api.semanticscholar.org/graph/v1/paper/search",
                params={
                    "query": en_query,
                    "limit": 5,
                    "fields": "title,authors,year,externalIds,url",
                },
                headers=headers,
            )
            for item in data.get("data") or []:
                external_ids = item.get("externalIds") or {}
                doi = external_ids.get("DOI") or ""
                paper_url = item.get("url") or ""
                papers.append(
                    {
                        "title": item.get("title") or "Sans titre",
                        "author": first_author_semantic(item),
                        "year": item.get("year") or "N/A",
                        "doi": doi,
                        "url": f"https://doi.org/{doi}" if doi else paper_url,
                        "source": "Semantic Scholar",
                    }
                )
        except Exception as exc:
            errors.append(f"Semantic Scholar: {exc}")

    # 4) Crossref
    if any("Crossref" in p for p in platforms):
        try:
            data = safe_get_json(
                "https://api.crossref.org/works",
                params={
                    "query.bibliographic": en_query,
                    "rows": 5,
                    "sort": "relevance",
                },
            )
            items = ((data.get("message") or {}).get("items")) or []
            for item in items:
                titles = item.get("title") or []
                title = titles[0] if titles else "Sans titre"
                doi = item.get("DOI") or ""
                papers.append(
                    {
                        "title": title,
                        "author": crossref_author(item),
                        "year": crossref_year(item),
                        "doi": doi,
                        "url": (
                            f"https://doi.org/{doi}"
                            if doi
                            else (item.get("URL") or "")
                        ),
                        "source": "Crossref",
                    }
                )
        except Exception as exc:
            errors.append(f"Crossref: {exc}")

    # 5) Europe PMC
    if any("PubMed" in p or "Europe PMC" in p for p in platforms):
        try:
            data = safe_get_json(
                "https://www.ebi.ac.uk/europepmc/webservices/rest/search",
                params={
                    "query": en_query,
                    "format": "json",
                    "pageSize": 5,
                    "resultType": "core",
                },
            )
            results = ((data.get("resultList") or {}).get("result")) or []
            for item in results:
                doi = item.get("doi") or ""
                pmid = item.get("pmid") or ""
                pmcid = item.get("pmcid") or ""

                if doi:
                    url = f"https://doi.org/{doi}"
                elif pmcid:
                    url = f"https://europepmc.org/article/PMC/{pmcid.replace('PMC', '')}"
                elif pmid:
                    url = f"https://europepmc.org/article/MED/{pmid}"
                else:
                    url = "https://europepmc.org/"

                papers.append(
                    {
                        "title": item.get("title") or "Sans titre",
                        "author": epmc_author(item),
                        "year": item.get("pubYear") or "N/A",
                        "doi": doi,
                        "url": url,
                        "source": "Europe PMC",
                    }
                )
        except Exception as exc:
            errors.append(f"Europe PMC: {exc}")

    # منصات تجارية: روابط بحث فقط ما لم يتوفر مفتاح/API رسمي مخصص
    if any("Elsevier" in p for p in platforms):
        portal_links.append(
            {
                "name": "ScienceDirect",
                "url": f"https://www.sciencedirect.com/search?qs={en_for_web}",
                "note": "بحث مباشر في ScienceDirect",
            }
        )

    if any("Emerald" in p for p in platforms):
        portal_links.append(
            {
                "name": "Emerald Insight",
                "url": f"https://www.emerald.com/insight/search?q={en_for_web}",
                "note": "بحث مباشر في Emerald Insight",
            }
        )

    if any("Taylor" in p for p in platforms):
        portal_links.append(
            {
                "name": "Taylor & Francis Online",
                "url": f"https://www.tandfonline.com/action/doSearch?AllField={en_for_web}",
                "note": "بحث مباشر في Taylor & Francis",
            }
        )

    if any("Clarivate" in p or "Scopus" in p for p in platforms):
        portal_links.append(
            {
                "name": "Scopus / Web of Science",
                "url": "https://www.scopus.com/search/form.uri?display=basic",
                "note": f"استخدم الاستعلام: {en_query}",
            }
        )

    papers = deduplicate_papers(papers)
    return papers[:12], portal_links, errors


def extract_uploaded_text(uploaded_file):
    if uploaded_file is None:
        return "", ""

    filename = uploaded_file.name
    try:
        if filename.lower().endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            text = " ".join((page.extract_text() or "") for page in reader.pages[:15])
        else:
            doc = Document(uploaded_file)
            text = " ".join(p.text for p in doc.paragraphs[:60])
        return text.strip(), ""
    except Exception as exc:
        return "", str(exc)


def make_safe_filename(value):
    value = re.sub(r'[\\/:*?"<>|]+', "_", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value[:70] or "research"



# ==========================================
# Gemini: Retry + Fallback anti-503
# ==========================================
def is_transient_gemini_error(error_text):
    error_lower = str(error_text).lower()
    transient_signals = (
        "503",
        "unavailable",
        "high demand",
        "overloaded",
        "temporarily unavailable",
        "429",
        "resource_exhausted",
        "deadline_exceeded",
        "timeout",
        "timed out",
    )
    return any(signal in error_lower for signal in transient_signals)


def call_gemini_with_fallback(api_key, prompt, sources_footer):
    """
    Essaie plusieurs modèles Gemini.
    Pour les erreurs temporaires 503/429:
      - nouvelle tentative avec backoff
      - passage au modèle suivant
    Une panne Gemini ne fait pas planter le reste de l'application.
    """
    client = genai.Client(api_key=api_key)

    # Modèles Flash actuels + solution de repli très stable.
    model_candidates = [
        "gemini-3.7-flash",
        "gemini-3.6-flash",
        "gemini-3.5-flash",
        "gemini-2.5-flash",
    ]

    attempts_per_model = 2
    error_history = []

    for model_name in model_candidates:
        for attempt in range(1, attempts_per_model + 1):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                )

                response_text = (getattr(response, "text", "") or "").strip()

                if not response_text:
                    raise RuntimeError(
                        f"Le modèle {model_name} a retourné une réponse vide."
                    )

                return {
                    "ok": True,
                    "model": model_name,
                    "text": (
                        f"### 🤖 Gemini model: `{model_name}`\n\n"
                        + response_text
                        + sources_footer
                    ),
                    "errors": error_history,
                }

            except Exception as exc:
                error_text = str(exc)
                error_history.append(
                    f"{model_name} - tentative {attempt}: {error_text}"
                )

                # Erreur non temporaire: clé invalide, permission, requête incorrecte...
                if not is_transient_gemini_error(error_text):
                    return {
                        "ok": False,
                        "model": model_name,
                        "text": (
                            "❌ خطأ Gemini غير مؤقت.\n\n"
                            f"`{error_text}`"
                        ),
                        "errors": error_history,
                    }

                # Une seule attente courte avant la deuxième tentative.
                if attempt < attempts_per_model:
                    wait_seconds = (2 ** attempt) + random.uniform(0.3, 1.2)
                    time.sleep(wait_seconds)
                else:
                    # Le modèle reste chargé: essayer le modèle suivant.
                    break

    last_error = (
        error_history[-1]
        if error_history
        else "Erreur Gemini inconnue."
    )

    return {
        "ok": False,
        "model": "",
        "text": (
            "⚠️ جميع نماذج Gemini المجربة غير متاحة مؤقتاً "
            "أو تحت ضغط مرتفع.\n\n"
            "التطبيق لم يتوقف: يمكنك الاعتماد على نتائج Groq "
            "والمحكم المنهجي، ثم إعادة تجربة Gemini لاحقاً.\n\n"
            f"آخر خطأ:\n`{last_error}`"
        ),
        "errors": error_history,
    }


# ==========================================
# 5. واجهة المستخدم الرئيسية
# ==========================================
st.title("🏛️ محطة عمل الباحث الذكي المتكاملة")
st.caption(
    "المنظومة الأكاديمية للبحث متعدد المصادر والتحكيم المقارن "
    "(ANRN Deep Research)"
)

with st.expander(
    "🔐 إعدادات المفاتيح السحابية (Groq / Gemini / Semantic Scholar)",
    expanded=False,
):
    st.caption(
        "الأفضل حفظ المفاتيح في Streamlit Secrets وعدم كتابتها داخل app.py."
    )
    col_k1, col_k2, col_k3 = st.columns(3)
    with col_k1:
        groq_key = st.text_input(
            "Groq API Key:",
            value=DEFAULT_GROQ_KEY,
            type="password",
        )
    with col_k2:
        gemini_key = st.text_input(
            "Google Gemini API Key:",
            value=DEFAULT_GEMINI_KEY,
            type="password",
        )
    with col_k3:
        s2_key = st.text_input(
            "Semantic Scholar API Key (اختياري):",
            value=DEFAULT_S2_KEY,
            type="password",
        )

with st.expander(
    "👤 بطاقة تعريف الباحث (توثيق بيانات صاحب البحث)",
    expanded=True,
):
    col_p1, col_p2 = st.columns(2)
    with col_p1:
        res_name = st.text_input("الاسم واللقب *", value="")
        res_role = st.selectbox(
            "الصفة الأكاديمية *",
            [
                "أستاذ جامعي / باحث دائم",
                "طالب دكتوراه",
                "طالب ماستر / تخرج",
                "باحث حر / مهني",
            ],
            index=0,
        )
        res_affil = st.text_input(
            "الانتماء المؤسسي / الجامعة / المخبر *",
            value="",
        )
    with col_p2:
        res_email = st.text_input("البريد الإلكتروني *", value="")
        res_phone = st.text_input("رقم الهاتف (اختياري)", value="")

st.markdown("### 🎛️ تخصيص موضوع ومنصات البحث")
col1, col2 = st.columns(2)
with col1:
    research_title = st.text_input(
        "عنوان البحث أو الإشكالية الرئيسية *:",
        key="form_title",
    )
with col2:
    research_field = st.text_input(
        "الميدان والتخصص الأكاديمي:",
        key="form_field",
    )

col_opt1, col_opt2 = st.columns(2)
with col_opt1:
    selected_language = st.selectbox(
        "🌐 لغة صياغة التقرير والمسودة الأكاديمية:",
        ["العربية", "English", "Français"],
        index=0,
    )
with col_opt2:
    selected_platforms = st.multiselect(
        "📚 اختر منصات البحث الأكاديمية المستهدفة:",
        options=[
            "ASJP (البوابة الجزائرية للمجلات العلمية)",
            "OpenAlex (الفهرس الشامل)",
            "Semantic Scholar (AI2)",
            "Crossref (توثيق DOI)",
            "PubMed / Europe PMC (العلوم الطبية والصحية)",
            "Elsevier (ScienceDirect)",
            "Emerald Insight (علوم التسيير والإدارة)",
            "Taylor & Francis (العلوم الإنسانية والاجتماعية)",
            "Clarivate / Scopus (الأوراق عالية الاقتباس)",
        ],
        default=[
            "ASJP (البوابة الجزائرية للمجلات العلمية)",
            "OpenAlex (الفهرس الشامل)",
            "Semantic Scholar (AI2)",
            "Crossref (توثيق DOI)",
        ],
    )

uploaded_file = st.file_uploader(
    "📁 أو اسحب وأفلت ملف بحثك (PDF / Word) للتحليل الهجين:",
    type=["pdf", "docx"],
    key="file_uploader",
)

extracted_text, upload_error = extract_uploaded_text(uploaded_file)
effective_title = research_title

if uploaded_file is not None:
    if upload_error:
        st.error(f"تعذر استخراج الملف: {upload_error}")
    else:
        st.success(f"✔ تم استخراج نصوص الملف: {uploaded_file.name} بنجاح.")
        clean_file_title = (
            uploaded_file.name.rsplit(".", 1)[0]
            .replace("_", " ")
            .replace("-", " ")
        )
        # لا نعدل session_state بعد إنشاء text_input حتى لا يظهر خطأ Streamlit.
        if not research_title.strip():
            effective_title = clean_file_title
            st.info(
                f"سيتم استخدام اسم الملف كعنوان مؤقت للبحث: {effective_title}"
            )

# التصحيح الأساسي للخطأ الظاهر في Streamlit:
# st.columns() غير صالح بدون spec.
col_btn1, col_btn2 = st.columns([3, 1])
with col_btn1:
    launch_btn = st.button(
        "🚀 بدء دورة البحث والتحكيم المقارن",
        type="primary",
    )
with col_btn2:
    st.button(
        "🧹 مسح وتفريغ",
        on_click=reset_all_fields,
        type="secondary",
    )


# ==========================================
# 6. إطلاق البحث والتحكيم
# ==========================================
if launch_btn:
    if not res_name.strip() or not res_affil.strip() or not res_email.strip():
        st.error(
            "⚠️ يرجى ملء الاسم، الانتماء المؤسسي والبريد الإلكتروني قبل بدء البحث."
        )
    elif not effective_title.strip():
        st.error("⚠️ يرجى إدخال عنوان البحث للمتابعة.")
    elif not selected_platforms:
        st.warning("⚠️ يرجى اختيار منصة بحث واحدة على الأقل.")
    elif not groq_key and not gemini_key:
        st.error(
            "⚠️ أضف مفتاح Groq أو Gemini على الأقل، ويفضل وضعه في Streamlit Secrets."
        )
    else:
        with st.spinner(
            "⏳ جاري استرجاع المصادر الأكاديمية ثم بناء التحليل المقارن..."
        ):
            papers_data, portal_links, search_errors = crawl_academic_papers(
                effective_title,
                selected_platforms,
                semantic_key=s2_key,
            )

            if search_errors:
                with st.expander("⚠️ تفاصيل أخطاء بعض مصادر البحث", expanded=False):
                    for err in search_errors:
                        st.write(f"- {err}")

            if portal_links:
                st.markdown("### 🔎 روابط البحث المباشر في المنصات")
                for link in portal_links:
                    st.markdown(
                        f"- **{link['name']}** — {link['note']}  \n"
                        f"  [فتح منصة البحث]({link['url']})"
                    )

            if papers_data:
                papers_summary = "\n".join(
                    (
                        f"- [{p['source']}] {p['title']} ({p['year']}) | "
                        f"المؤلف: {p['author']} | DOI: {p['doi'] or 'N/A'} | "
                        f"الرابط: {p['url']}"
                    )
                    for p in papers_data
                )

                bibtex_items = []
                for i, p in enumerate(papers_data):
                    bibtex_items.append(
                        "@article{"
                        f"ref{i + 1}_{p['year']},\n"
                        f"  title={{{p['title']}}},\n"
                        f"  author={{{p['author']}}},\n"
                        f"  year={{{p['year']}}},\n"
                        f"  doi={{{p['doi']}}},\n"
                        f"  url={{{p['url']}}}\n"
                        "}"
                    )
                bibtex_text = "\n\n".join(bibtex_items)

                sources_footer = (
                    "\n\n---\n"
                    "### 📚 المراجع الأكاديمية المسترجعة وروابط التحقق:\n"
                    + "\n".join(
                        (
                            f"* **{p['title']}** ({p['year']}) — "
                            f"*{p['author']}* — {p['source']}  \n"
                            f"  [رابط التحقق]({p['url']})"
                            + (f" | DOI: `{p['doi']}`" if p["doi"] else "")
                        )
                        for p in papers_data
                    )
                )
            else:
                papers_summary = (
                    "لم تُسترجع مقالات منظمة عبر APIs في هذه الدورة. "
                    "لا تنشئ أو تختلق مراجع أو DOI. "
                    "يمكن تقديم إطار بحثي عام فقط، مع التصريح بغياب المصادر المسترجعة."
                )
                bibtex_text = ""
                sources_footer = (
                    "\n\n---\n"
                    "### 📚 ملاحظة المصادر\n"
                    "لم تُسترجع مراجع منظمة عبر APIs في هذه الدورة. "
                    "يرجى استخدام روابط المنصات أعلاه لإتمام التحقق اليدوي."
                )

            uploaded_context = ""
            if extracted_text:
                uploaded_context = (
                    "\n\nمقتطف من الملف الذي رفعه الباحث:\n"
                    + extracted_text[:14000]
                )

            if selected_language == "English":
                lang_instruction = (
                    "Write the whole analysis in formal academic English."
                )
            elif selected_language == "Français":
                lang_instruction = (
                    "Rédiger toute l'analyse en français académique rigoureux."
                )
            else:
                lang_instruction = (
                    "اكتب التحليل كاملاً باللغة العربية الأكاديمية الرصينة."
                )

            prompt = f"""
أنت خبير في التحكيم الأكاديمي واكتشاف الفجوات العلمية
وفق مصفوفة الفجوات السباعية (7D Gap Taxonomy).

{lang_instruction}

موضوع البحث: {effective_title}
الميدان: {research_field}
الباحث: {res_name} ({res_role} - {res_affil})

المصادر الأكاديمية المسترجعة:
{papers_summary}
{uploaded_context}

قواعد إلزامية:
- لا تختلق أي مرجع أو DOI أو اسم مؤلف.
- استخدم فقط المراجع الموجودة فعلاً في قائمة المصادر المسترجعة عند الاستشهاد.
- إذا كانت المصادر غير كافية، صرّح بذلك بوضوح.
- ميّز بين ما هو مستند إلى المصادر وما هو اقتراح منهجي منك.

المطلوب:
1. مصفوفة الإطباق المنهجي (Methodological Overlap Matrix).
2. استخراج 3 فجوات بحثية نوعية ومخصصة لموضوع البحث.
3. صياغة إشكالية بحث دقيقة.
4. صياغة 3 فرضيات علمية قابلة للاختبار.
5. اقتراح المنهجية وأدوات القياس والتحليل المناسبة.
6. إعداد مسودة أكاديمية تشمل مقدمة، إطاراً نظرياً مختصراً،
   المنهجية المقترحة والنتائج المتوقعة وحدود الدراسة.
""".strip()

            groq_output = ""
            gemini_output = ""
            method_output = ""

            # Groq - Reviewer 1
            if groq_key:
                try:
                    groq_client = Groq(api_key=groq_key)
                    groq_res = groq_client.chat.completions.create(
                        model="openai/gpt-oss-20b",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.3,
                    )
                    groq_output = (
                        (groq_res.choices[0].message.content or "")
                        + sources_footer
                    )
                except Exception as exc:
                    groq_output = f"خطأ في مسار Groq: {exc}"
            else:
                groq_output = "لم يتم إدخال مفتاح Groq."

            # Gemini - Reviewer 2 avec Retry + Fallback anti-503
            if gemini_key:
                gemini_result = call_gemini_with_fallback(
                    gemini_key,
                    prompt,
                    sources_footer,
                )
                gemini_output = gemini_result["text"]
                gemini_model_used = gemini_result["model"]
                gemini_errors = gemini_result["errors"]
            else:
                gemini_output = "لم يتم إدخال مفتاح Gemini."
                gemini_model_used = ""
                gemini_errors = []

            # Groq - Reviewer 3 (منهجي/إحصائي)
            if groq_key:
                try:
                    method_prompt = (
                        "أنت محكّم منهجي وإحصائي متخصص في تصميم البحوث "
                        "وبناء نماذج القياس وPLS-SEM عند ملاءمتها. "
                        "لا تفترض أن PLS-SEM مناسب دائماً؛ برر اختيارك. \n\n"
                        + prompt
                    )
                    method_res = groq_client.chat.completions.create(
                        model="openai/gpt-oss-20b",
                        messages=[{"role": "user", "content": method_prompt}],
                        temperature=0.25,
                    )
                    method_output = (
                        (method_res.choices[0].message.content or "")
                        + sources_footer
                    )
                except Exception as exc:
                    method_output = f"خطأ في مسار المحكم المنهجي: {exc}"
            else:
                method_output = "لم يتم إدخال مفتاح Groq."

            st.session_state["results"] = {
                "title": effective_title,
                "field": research_field,
                "language": selected_language,
                "researcher": {
                    "name": res_name,
                    "role": res_role,
                    "affil": res_affil,
                    "email": res_email,
                    "phone": res_phone,
                },
                "papers": papers_data,
                "portal_links": portal_links,
                "bibtex": bibtex_text,
                "groq": groq_output,
                "gemini": gemini_output,
                "gemini_model": gemini_model_used,
                "gemini_errors": gemini_errors,
                "method": method_output,
            }

            log_id = save_research_log(
                {
                    "name": res_name,
                    "role": res_role,
                    "affiliation": res_affil,
                    "email": res_email,
                    "phone": res_phone,
                    "title": effective_title,
                    "field": research_field,
                    "language": selected_language,
                }
            )
            st.session_state["current_log_id"] = log_id
            st.success("🎉 اكتملت دورة البحث والتحكيم المقارن.")


# ==========================================
# 7. عرض النتائج والتحميل
# ==========================================
if "results" in st.session_state:
    res = st.session_state["results"]

    st.markdown("---")
    st.header(
        f"📝 مخرجات التحكيم المقارن ومسودات البحث ({res['language']})"
    )
    st.caption(
        f"👤 الباحث: **{res['researcher']['name']}** "
        f"({res['researcher']['role']} - {res['researcher']['affil']})"
    )

    if res.get("gemini_model"):
        st.success(
            "🤖 Gemini يعمل حالياً باستخدام النموذج: "
            f"{res['gemini_model']}"
        )

    if res.get("gemini_errors"):
        with st.expander(
            "🧪 تفاصيل محاولات Gemini السابقة",
            expanded=False,
        ):
            for gemini_error in res["gemini_errors"]:
                st.code(gemini_error)

    if res["papers"]:
        st.markdown("### 📚 المقالات المسترجعة آلياً")
        papers_df = pd.DataFrame(res["papers"])[
            ["source", "title", "author", "year", "doi", "url"]
        ]
        st.dataframe(papers_df, use_container_width=True, hide_index=True)
    else:
        st.warning(
            "لم تُسترجع مقالات منظمة آلياً. "
            "النتائج النصية يجب اعتبارها إطاراً بحثياً يحتاج تحققاً مرجعياً."
        )

    tab1, tab2, tab3, tab4 = st.tabs(
        [
            "🔹 Google Gemini",
            "🔹 المحكم المنهجي",
            "🔹 Groq",
            "📚 BibTeX",
        ]
    )

    with tab1:
        st.markdown(res["gemini"])
    with tab2:
        st.markdown(res["method"])
    with tab3:
        st.markdown(res["groq"])
    with tab4:
        if res["bibtex"]:
            st.code(res["bibtex"], language="latex")
        else:
            st.info("لا توجد مراجع BibTeX في هذه الدورة.")

    doc = Document()
    doc.add_heading(
        f"Academic Research & Multi-Model Review: {res['title']}",
        0,
    )
    doc.add_paragraph(
        f"Researcher: {res['researcher']['name']} "
        f"({res['researcher']['role']} - {res['researcher']['affil']})"
    )
    doc.add_paragraph(
        f"Field: {res['field']} | Language: {res['language']}"
    )

    doc.add_heading("1. Google Gemini Review", level=1)
    doc.add_paragraph(res["gemini"])
    doc.add_heading("2. Methodological Reviewer", level=1)
    doc.add_paragraph(res["method"])
    doc.add_heading("3. Groq Review", level=1)
    doc.add_paragraph(res["groq"])
    doc.add_heading("4. References & BibTeX", level=1)
    doc.add_paragraph(res["bibtex"] or "No structured references retrieved.")

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    safe_base = make_safe_filename(res["title"])
    col_d1, col_d2 = st.columns(2)

    with col_d1:
        st.download_button(
            label="📥 تحميل الورقة البحثية كملف Word (.docx)",
            data=doc_io.getvalue(),
            file_name=f"{safe_base}-{res['language']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
        )

    with col_d2:
        st.download_button(
            label="📚 تحميل مراجع BibTeX (.bib)",
            data=res["bibtex"] or "% No structured references retrieved.\n",
            file_name="references.bib",
            mime="text/plain",
        )

    st.markdown("---")
    st.markdown("### 💬 شاركنا بملاحظاتك المنهجية")
    feedback_input = st.text_area(
        "أدخل ملاحظاتك أو مقترحاتك:",
        placeholder="اكتب اقتراحك هنا...",
    )

    if st.button("📤 إرسال الاقتراح للمطور"):
        if not feedback_input.strip():
            st.warning("يرجى كتابة نص الاقتراح قبل الإرسال.")
        elif not st.session_state.get("current_log_id"):
            st.warning("لا يوجد سجل بحث حالي لربط الملاحظة به.")
        elif update_feedback(
            st.session_state["current_log_id"],
            feedback_input.strip(),
        ):
            st.success("✔ تم حفظ اقتراحك بنجاح.")


# ==========================================
# 8. بوابة المطور
# ==========================================
st.markdown("---")
with st.expander(
    "🛠️ بوابة المطور وسجل الباحثين والمقترحات",
    expanded=False,
):
    if not ADMIN_PIN:
        st.info(
            "بوابة المطور معطلة حالياً. "
            "أضف ADMIN_PIN داخل Streamlit Secrets لتفعيلها بأمان."
        )
    else:
        dev_pin = st.text_input(
            "أدخل رمز مرور المطور:",
            type="password",
            key="dev_pass",
        )

        if dev_pin == ADMIN_PIN:
            st.success("🔓 تم الدخول إلى سجلات المطور.")
            with get_db_connection() as conn:
                df_logs = pd.read_sql_query(
                    """
                    SELECT id, timestamp, researcher_name, role, affiliation,
                           email, phone, title, field, language, feedback
                    FROM research_logs
                    ORDER BY id DESC
                    """,
                    conn,
                )

            st.metric(
                "📊 إجمالي الأبحاث المسجلة في المنصة:",
                len(df_logs),
            )
            st.dataframe(
                df_logs,
                use_container_width=True,
                hide_index=True,
            )

            csv_data = df_logs.to_csv(index=False).encode("utf-8-sig")
            st.download_button(
                label="📥 تنزيل سجل الباحثين والمقترحات (CSV)",
                data=csv_data,
                file_name=(
                    "researchers_database_"
                    f"{datetime.now().strftime('%Y%m%d')}.csv"
                ),
                mime="text/csv",
            )

        elif dev_pin:
            st.error("❌ رمز المرور غير صحيح.")
